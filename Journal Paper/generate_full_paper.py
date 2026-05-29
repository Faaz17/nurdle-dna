"""
Generates a full IEEE 2-column journal paper for NurdleDNA.

Combines teammates' Introduction, Literature Review, and Methodology
(transcribed and tightened from the prior single-column submission)
with the Results / Discussion / Conclusion sections, in a layout
designed to fit ~6 pages.

Figure placeholders are inserted for the four images that are owned
by the team (CAD photo, 4-panel detection frames, FSM state diagram,
dashboard screenshot) — drop those in directly in Word.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.normpath(os.path.join(HERE, "..", "Paper Figures"))
OUT_PATH = os.path.join(HERE, "NurdleDNA_Full_Paper.docx")

# =====================================================================
# Helpers
# =====================================================================
def set_section_columns(section, n_cols, space_twips=420):
    """Set the column count for a section via direct sectPr XML edit."""
    sectPr = section._sectPr
    for cols in sectPr.findall(qn("w:cols")):
        sectPr.remove(cols)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(n_cols))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)

def add_section_break(doc, n_cols):
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.625)
    sec.right_margin = Inches(0.625)
    set_section_columns(sec, n_cols)
    return sec

def add_para(text, *, justify=True, indent=True, size=10,
             bold=False, italic=False, align=None,
             space_before=0, space_after=2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_section_heading(text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = True
    return p

def add_subsection_heading(text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = True
    r.italic = True
    return p

def add_figure(filename, caption, width_inches=3.3):
    img_path = os.path.join(FIG_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        r = p.add_run(f"[ image not found: {filename} ]")
        r.italic = True
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cr = cap.add_run(caption)
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(8.5)
    cr.italic = True

def add_wide_figure(filename, caption, width_inches=7.0):
    """Span both columns: 2-col → 1-col → figure → 2-col."""
    add_section_break(doc, 1)
    add_figure(filename, caption, width_inches=width_inches)
    add_section_break(doc, 2)

def add_figure_placeholder(label, what, caption, wide=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "B0B0B0")
        pBdr.append(b)
    pPr.append(pBdr)
    r1 = p.add_run(f"[ INSERT {label} ]\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    r1.font.size = Pt(9)
    r2 = p.add_run(what)
    r2.italic = True
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cr = cap.add_run(caption)
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(8.5)
    cr.italic = True

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)

# =====================================================================
# Document setup
# =====================================================================
doc = Document()

sec1 = doc.sections[0]
sec1.page_width = Inches(8.5)
sec1.page_height = Inches(11)
sec1.top_margin = Inches(0.75)
sec1.bottom_margin = Inches(1.0)
sec1.left_margin = Inches(0.625)
sec1.right_margin = Inches(0.625)
set_section_columns(sec1, 1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.08
style.paragraph_format.space_after = Pt(2)

# =====================================================================
# Title block (single column)
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(6)
tr = title.add_run(
    "NurdleDNA: An Edge-AI and IoT Inline System "
    "for Source-Level Microplastic Detection"
)
tr.bold = True
tr.font.name = "Times New Roman"
tr.font.size = Pt(20)

# Author block — 3 columns × 2 rows
authors = [
    ("Faaz Ali Sayyed",         "fasas305@uowmail.edu.au"),
    ("Muhammad Haaziq",         "mh084@uowmail.edu.au"),
    ("Daniel Joseph Koshy",     "djk545@uowmail.edu.au"),
    ("Mohammed Abdul Rehman",   "mar699@uowmail.edu.au"),
    ("Syed Hasan Ali",          "sha245@uowmail.edu.au"),
]

author_table = doc.add_table(rows=2, cols=3)
author_table.autofit = True

def fill_author_cell(cell, name, email):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(2)
    lines = [
        (name,                                        True,  False, 10.5),
        ("Computer and Autonomous Systems Engineering", False, False, 9.5),
        ("University of Wollongong in Dubai",          False, False, 9.5),
        ("Dubai, United Arab Emirates",                False, False, 9.5),
        (email,                                        False, True,  9.5),
    ]
    for i, (txt, bold, italic, sz) in enumerate(lines):
        if i > 0:
            para.add_run("\n")
        r = para.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(sz)
        r.bold = bold
        r.italic = italic

for i, (name, email) in enumerate(authors):
    row = i // 3
    col = i % 3
    fill_author_cell(author_table.rows[row].cells[col], name, email)

remove_table_borders(author_table)

# Switch to 2-column body
add_section_break(doc, 2)

# =====================================================================
# Abstract + Keywords
# =====================================================================
abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_p.paragraph_format.space_before = Pt(4)
abs_p.paragraph_format.space_after = Pt(4)

ah = abs_p.add_run("Abstract")
ah.bold = True; ah.italic = True
ah.font.name = "Times New Roman"; ah.font.size = Pt(9.5)
dash = abs_p.add_run("—")
dash.bold = True; dash.font.size = Pt(9.5)
ab = abs_p.add_run(
    "Industrial coastal water infrastructure requires rapid pollution "
    "interception before microplastic contamination leaves controlled "
    "discharge pathways and enters marine ecosystems. Pre-production "
    "plastic pellets (nurdles) are difficult to manage because they are "
    "small, highly mobile, visually similar to other suspended particles, "
    "and handled in large quantities across petrochemical plants, ports, "
    "and industrial drainage networks. Conventional approaches — manual "
    "inspection, passive filtration, turbidity sensing, and laboratory "
    "FTIR or Raman analysis — either lack real-time capability or fail "
    "to provide source-level containment. This paper introduces "
    "NurdleDNA, an edge-AI and IoT-enabled inline system for real-time "
    "microplastic detection, automated flow isolation, physical evidence "
    "capture, and cloud-based compliance logging. The architecture "
    "integrates an optical vision chamber, YOLO-based detection, "
    "turbidity verification, VOC/gas sensing, load-cell capture "
    "measurement, servo-actuated valve isolation, and a cloud audit "
    "dashboard within a closed-loop detect–classify–actuate–capture–"
    "report framework. Prototype-level validation demonstrates rapid "
    "contamination recognition, automatic valve response, evidence "
    "capture, and time-stamped event logging for scalable deployment in "
    "industrial water channels."
)
ab.italic = True; ab.bold = True
ab.font.name = "Times New Roman"; ab.font.size = Pt(9.5)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.paragraph_format.space_after = Pt(4)
kh = kw_p.add_run("Keywords")
kh.bold = True; kh.italic = True
kh.font.name = "Times New Roman"; kh.font.size = Pt(9.5)
kw_p.add_run("—").bold = True
kt = kw_p.add_run(
    "microplastic detection, nurdles, edge AI, Internet of Things, "
    "industrial water monitoring, source-level containment, multi-sensor "
    "fusion, cloud audit logging."
)
kt.italic = True
kt.font.name = "Times New Roman"; kt.font.size = Pt(9.5)

# =====================================================================
# I. Introduction
# =====================================================================
add_section_heading("I.  Introduction")
add_para(
    "Plastic pellet pollution, specifically from pre-production pellets "
    "known as nurdles, has emerged as a critical global threat to marine "
    "and coastal ecosystems. Due to their small size (3–5 mm) and high "
    "mobility, these pellets are easily transported by winds and ocean "
    "currents, accumulating in ecological hotspots and acting as vectors "
    "for toxic chemical contaminants when ingested by marine life [9]. "
    "Tracing these pollutants back to their industrial sources remains a "
    "major challenge due to their apparent physical homogeneity."
)
add_para(
    "To resolve this, we propose the concept of “Nurdle DNA”, "
    "a multi-dimensional physical and chemical fingerprint unique to "
    "individual pellet batches and their environmental histories. "
    "Decoding this signature requires a combination of high-throughput "
    "optical and spectroscopic techniques: portable Raman spectroscopy "
    "for primary polymer identification and surface weathering [9], "
    "visible–near-infrared (VIS-NIR) spectroscopy combined with "
    "chemometric machine learning such as 1D-CNNs and partial least "
    "squares regression for complex mixtures [10], and state-of-the-art "
    "computer vision object detection (YOLOv8, RetinaNet, Faster R-CNN) "
    "for automated physical classification at industrial scale [11]. "
    "Integrated, these techniques establish a framework for tracking "
    "spill origins, assessing marine residence times, and optimising "
    "coastal monitoring."
)

# =====================================================================
# II. Literature Review
# =====================================================================
add_section_heading("II.  Literature Review")

add_subsection_heading("A.  Microplastic and Nurdle Pollution in Industrial Water Systems")
add_para(
    "Microplastics — plastic particles smaller than 5 mm — persist in "
    "water systems, fragment further, and may enter food chains [1]. "
    "Studies from Dubai and the wider MENA region confirm the need for "
    "coastal and industrial microplastic monitoring [2], [3]. "
    "Pre-production nurdles are a major source-level risk because they "
    "leak from polymer plants, ports, and logistics areas into drains "
    "and coastal waters. NurdleDNA addresses this gap through an inline "
    "IoT system that detects nurdles, isolates contaminated flow, "
    "captures evidence, and logs events to the cloud."
)

add_subsection_heading("B.  Climate, Ecosystem, and Environmental Damage")
add_para(
    "Microplastic pollution affects ecosystems, biodiversity, and "
    "climate resilience. UNEP reports that plastic pollution alters "
    "habitats and natural processes, reducing ecosystem resilience to "
    "climate change [4]. In marine environments, microplastics affect "
    "ocean health and reduce the ocean’s capacity to support "
    "carbon sequestration, which is important for climate regulation "
    "[5]. Preventing nurdle leakage at the industrial source is "
    "therefore more effective than recovering plastics after "
    "environmental dispersion."
)

add_subsection_heading("C.  Limitations of Existing Monitoring Approaches")
add_para(
    "Current detection methods are limited for real-time industrial "
    "prevention. Manual inspection is slow, filters can clog, and "
    "laboratory methods such as FTIR, Raman, and microscopy require "
    "offline sampling and specialised equipment [6]. Turbidity sensors "
    "and particle counters can detect abnormalities but cannot reliably "
    "classify nurdles from bubbles, sand, or biological particles. This "
    "creates a clear need for a system that can detect and stop "
    "contamination in real time."
)

add_subsection_heading("D.  Edge-AI, Multi-Sensor Fusion, and IoT Monitoring")
add_para(
    "Edge AI and IoT sensing provide a practical pathway for real-time "
    "microplastic monitoring. YOLO-based vision detects objects with "
    "bounding boxes, class labels, and confidence scores [7], while "
    "NVIDIA Jetson platforms support local AI inference without full "
    "cloud dependence [8]. NurdleDNA combines AI vision, turbidity "
    "sensing, VOC/gas detection, load-cell measurement, Arduino "
    "control, and cloud logging to support reliable multi-sensor "
    "environmental monitoring."
)

# =====================================================================
# III. Methodology
# =====================================================================
add_section_heading("III.  Methodology")

add_subsection_heading("A.  System Architecture and Physical Partitioning")
add_para(
    "The NurdleDNA prototype uses a dual-layer containment model "
    "designed in Autodesk Fusion 360 [12], isolating liquid-conveying "
    "pathways from processing electronics. The "
)
# Continue inline by appending a new para that doesn't indent (run-on)
add_para(
    "wet path (sensor zone) comprises a transparent inline optical flow "
    "cell with a WS2812B-8 LED strip providing uniform backlighting and "
    "an opposing camera viewport. A light-dependent resistor (LDR) "
    "inside the optical path serves as a turbidity sensor. An MQ-135 gas "
    "sensor in the headspace detects volatile organic compounds (VOCs). "
    "A downstream evidence cartridge is suspended from a load cell, and "
    "an upstream MG996R servo-actuated pinch valve governs flow "
    "isolation. The dry path (electronics zone) houses an NVIDIA Jetson "
    "Nano hosting the CNN inference engine and cloud publisher [13], "
    "and an Arduino Uno running the deterministic control logic. The "
    "two layers are physically partitioned so that water leaks cannot "
    "reach electronics and so a technician can swap sensors without "
    "touching the compute hardware.",
    indent=True,
)

add_figure_placeholder(
    "FIG. 1 — CAD model",
    "Fusion 360 isometric of the NurdleDNA prototype, showing the "
    "electronics enclosure mounted above the inline cylindrical flow "
    "conduit.",
    "Fig. 1. CAD model of the NurdleDNA prototype showing the electronics "
    "enclosure mounted above the inline cylindrical flow conduit."
)

add_subsection_heading("B.  Multi-Sensor Acquisition and Signal Conditioning")
add_para(
    "Signal conditioning and telemetry are handled by the Arduino Uno "
    "at a 5 Hz sampling rate. Optoelectronic turbidity is read on pin "
    "A0; values exceeding "
    "Vₒ > 600 (warning threshold) flag potential contamination. "
    "Volatile petrochemical hydrocarbons are monitored on pin A1 via "
    "the MQ-135; ADC values above 200 trigger a critical alarm. "
    "Gravimetric accumulation in the downstream cartridge is measured "
    "using a parallel-beam load cell interfaced with a 24-bit HX711 "
    "converter (pins D4 and D10), calibrated with set_scale() to "
    "resolve spill mass in grams."
)

add_subsection_heading("C.  Edge Computer Vision and Classification Pipeline")
add_para(
    "Real-time visual monitoring uses a hybrid pipeline on the Jetson "
    "Nano balancing deep-learning classification with classical safety "
    "nets [12]. A YOLOv8n model (2.6 M parameters) was trained for 25 "
    "epochs on a T4 GPU using the Roboflow microplastics-t0ddd v6 "
    "dataset (3,102 images, 19 classes) [14]. The trained model was "
    "exported to ONNX and executed via onnxruntime, reducing runtime "
    "memory from 1.0 GB to ~100 MB while sustaining >5 FPS on the "
    "Jetson Nano CPU. In parallel, a classical HSV colour-thresholding "
    "and contour-analysis algorithm runs on the central frame as a "
    "safety net for bright rounded geometries that fall below the "
    "YOLO confidence threshold."
)
add_para(
    "Four filtering layers suppress false positives: (1) a "
    "region-of-interest mask restricting detection to the central 60% "
    "of the frame; (2) a circularity filter ("
    "4πA/P²"
    ") rejecting linear highlights; (3) an exponential moving average "
    "(α = 0.20) smoothing per-frame counts; and (4) confirmation "
    "timers requiring warning (S2) and alarm (S3) states to persist for "
    "1.5 s and 2.5 s respectively before transitioning."
)

add_subsection_heading("D.  Deterministic Finite State Machine Control Logic")
add_para(
    "Industrial safety mandates deterministic actuation decoupled from "
    "probabilistic AI inference. Control is managed by a 5-state Moore "
    "FSM executed on the Arduino Uno [15]. S0 (INIT) is the boot state "
    "and holds the pinch valve OPEN. S1 (SysOK) is normal operation "
    "(green LED, valve OPEN). S2 (Caution/WARN) is triggered by "
    "LDR > 600 OR AI detection count ≥ 4 sustained for ≥ 1.5 "
    "s (yellow LED). S3 (ALARM/CRIT) is triggered by MQ-135 > 200 OR "
    "AI count ≥ 12 sustained for ≥ 2.5 s — the red LED and "
    "buzzer activate, and the servo pinch valve closes within ~300 ms. "
    "S3 is physically latched and is cleared only via S4 (RSTIN) "
    "through the reset button on pin D2. Telemetry between the Jetson "
    "Nano and Arduino Uno is exchanged at 5 Hz over a 115,200 baud USB "
    "serial link in JSON format."
)

add_subsection_heading("E.  Cloud Telemetry and Remote Monitoring Interface")
add_para(
    "Remote monitoring is supported via Firebase Realtime Database. "
    "The Jetson Nano streams sensor telemetry at 5 Hz; video frames "
    "are downsampled to 320×240 pixels, JPEG-encoded at 60% "
    "quality, base64-encoded, and uploaded at 1 Hz to the operator web "
    "dashboard. On entering the S3 ALARM state, an immutable audit "
    "event containing the UTC timestamp, sensor values, and per-class "
    "detection counts is pushed to the database to establish a "
    "regulatory compliance log."
)

# =====================================================================
# IV. Results and Discussion
# =====================================================================
add_section_heading("IV.  Results and Discussion")
add_para(
    "Each stage of the closed-loop detect–classify–actuate–capture–"
    "report pipeline is evaluated below. The overall system "
    "architecture and physical partitioning are summarised in Fig. 2."
)

add_figure(
    "fig1_block_diagram.png",
    "Fig. 2. System architecture showing the wet path (inline water flow "
    "with optical, gas, and gravimetric sensing) and the dry path "
    "(NVIDIA Jetson Nano running the YOLOv8n + OpenCV pipeline, Arduino "
    "Uno running the FSM, and the Firebase-backed audit dashboard).",
    width_inches=3.3,
)

add_subsection_heading("A.  Detection Accuracy of the Edge Vision Pipeline")
add_para(
    "The YOLOv8n model achieved mean Average Precision at IoU 0.50 "
    "(mAP@50) of 0.81 for synthetic pens, 0.75 for air bubbles, and "
    "0.60 for fragments. The higher precision on pens and bubbles "
    "reflects their consistent geometric structure; fragments suffer "
    "from intra-class shape variability and visual overlap with "
    "biological debris. This skew is consistent with prior reports on "
    "YOLO-family models applied to microplastic imagery [12]. The "
    "bubble class result is operationally important: bubbles are the "
    "dominant source of optical false positives in inline flow cells, "
    "and reliable bubble labelling enables downstream filtering rather "
    "than confusion with true nurdles. The classical OpenCV "
    "HSV–contour safety net contributes redundancy when the deep model "
    "falls below threshold."
)

add_figure(
    "fig5_training_curves.png",
    "Fig. 3. YOLOv8n training loss and validation metrics over 25 "
    "epochs on the Roboflow microplastics-t0ddd v6 dataset; final "
    "mAP@50 ≈ 0.72.",
    width_inches=3.3,
)
add_figure(
    "fig6_confusion_matrix.png",
    "Fig. 4. Normalised confusion matrix (consolidated 8-class view) "
    "showing strong diagonal recovery for nurdle (0.78), pen (0.81), "
    "and bubble (0.75); fragment shows the largest off-diagonal spread.",
    width_inches=3.3,
)
add_figure_placeholder(
    "FIG. 5 — Detection frames (4-panel)",
    "Four-panel image: (a) clean baseline; (b) bubble correctly "
    "labelled; (c) confirmed nurdle; (d) edge case where the OpenCV "
    "safety net fires below YOLO threshold.",
    "Fig. 5. Representative detection outputs from the hybrid YOLOv8n + "
    "OpenCV pipeline."
)

add_subsection_heading("B.  Edge Inference Efficiency")
add_para(
    "Exporting the model to ONNX and executing via onnxruntime reduced "
    "the runtime memory footprint by an order of magnitude (1.0 GB to "
    "~100 MB), essential given the Jetson Nano’s 4 GB shared "
    "memory budget. CPU inference sustained >5 FPS on 320×240 "
    "input, exceeding the 5 Hz control-loop requirement. End-to-end "
    "response latency from pellet entry to dashboard update was "
    "approximately 3.0 s, aggregating frame capture, ONNX inference, "
    "the temporal-filter confirmation window, USB-serial transit, and "
    "the Firebase publish. The 2.5 s S3 confirmation window dominates "
    "the budget (Fig. 6)."
)
add_figure(
    "fig4_latency_breakdown.png",
    "Fig. 6. End-to-end latency breakdown. The temporal-filter "
    "confirmation window (2.5 s) dominates; inference and network "
    "transit contribute only ~430 ms in aggregate.",
    width_inches=3.3,
)

add_subsection_heading("C.  Multi-Sensor Fusion and False-Positive Suppression")
add_para(
    "The four temporal-filter layers (ROI mask, circularity, EMA, and "
    "confirmation timers) jointly suppressed transient optical "
    "artefacts during bench testing. Across a 30-minute trial with "
    "deliberate disturbances (air injection, ambient-light flicker, "
    "turbid water), the system produced three spurious S2 transitions "
    "and no spurious S3 transitions before reaching steady state; no "
    "S3 ALARM was triggered by air bubbles alone in any test run."
)
add_para(
    "Multi-sensor fusion is the core actuation property of NurdleDNA. "
    "S2 is entered on LDR > 600 OR AI count ≥ 4 sustained ≥ "
    "1.5 s; S3 on MQ-135 > 200 OR AI count ≥ 12 sustained "
    "≥ 2.5 s. The OR-based escalation provides graceful "
    "degradation: if the camera is occluded or lighting fails, the "
    "analog sensors still drive the FSM through the contamination "
    "pathway; if the gas sensor saturates, the vision pipeline retains "
    "containment authority. Confirmation windows absorb the higher "
    "per-channel false-positive rate that this redundancy implies."
)

add_subsection_heading("D.  Deterministic Actuation and Containment Performance")
add_para(
    "The MG996R servo-actuated pinch valve closed within ~300 ms of "
    "FSM entry to S3 (command-to-seating). Because the FSM uses "
    "deterministic threshold logic, the actuation path is independent "
    "of Jetson Nano availability. During a deliberate USB-serial "
    "disconnection test, the Arduino independently entered S3 within "
    "~500 ms of an LDR + MQ-135 co-excursion (one 5 Hz sample period "
    "≤ 200 ms + instantaneous MQ-135 trigger + ~300 ms valve "
    "seating), confirming the intended fail-operational behaviour. "
    "S3 is physically latched and is cleared only via the operator "
    "reset on D2 — consistent with industrial safety practice for "
    "closed-loop fluid containment, as automatic re-opening on a "
    "transient sensor clearance would risk releasing accumulated "
    "contamination."
)

add_figure(
    "fig9_timing_trace.png",
    "Fig. 7. Time-aligned single-event trace. LDR exceeds the WARN "
    "threshold at t ≈ 3.7 s; after the 1.5 s confirmation the FSM "
    "enters S2 (amber). AI count crosses CRIT at t ≈ 5.2 s; after "
    "the 2.5 s confirmation the FSM enters S3 (red) and the servo "
    "transitions from 90° to 0° within ~300 ms.",
    width_inches=3.3,
)
add_figure_placeholder(
    "FIG. 8 — FSM state diagram",
    "Five-state Moore machine: S0 INIT → S1 SysOK → S2 WARN "
    "→ S3 ALARM (latched) → S4 RSTIN, with transition "
    "triggers and per-state outputs.",
    "Fig. 8. Five-state Moore FSM governing actuation. S3 ALARM is "
    "physically latched and is cleared only by operator reset on D2 "
    "(transition through S4)."
)

add_subsection_heading("E.  Evidence Capture and Gravimetric Logging")
add_para(
    "The HX711-interfaced load cell, calibrated via set_scale() against "
    "known reference masses, resolved cartridge accumulation in grams "
    "at the Arduino’s 5 Hz acquisition rate (Fig. 9). During the "
    "AURAK live demonstration on 21 May 2026, the captured-mass channel "
    "recorded approximately 2.4 g of accumulated material over the "
    "live run, rendered on the dashboard in real time alongside the AI "
    "density index. The combination of visual evidence (the rendered "
    "detection frame), gravimetric evidence (load cell), and chemical "
    "evidence (MQ-135 VOC excursion) constitutes a three-channel "
    "forensic record per alarm event — a property difficult to achieve "
    "with single-modality systems."
)
add_figure(
    "fig10_calibration.png",
    "Fig. 9. Load-cell calibration. Reference masses 0–200 g produce a "
    "linear response with R² ≈ 1.000 and sensitivity ~412 "
    "counts/g (~2.4 mg per count).",
    width_inches=3.3,
)

add_subsection_heading("F.  Cloud Telemetry and Audit Logging")
add_para(
    "Firebase Realtime Database integration sustained 5 Hz sensor "
    "telemetry and a 1 Hz visual stream (320×240 JPEG at 60% "
    "quality, base64-encoded). The decoupled rates kept total upstream "
    "bandwidth to ~200 kbps — comfortably within the budget of an "
    "LTE-class industrial gateway — dominated by the 1 Hz visual "
    "channel (~15 KB per JPEG, inflated ~33% by base64). On every S3 "
    "ALARM transition, an immutable audit event containing the UTC "
    "timestamp, sensor values, and per-class detection counts was "
    "committed to the database; this event log is the regulatory "
    "artefact that distinguishes NurdleDNA from passive monitoring."
)
add_figure_placeholder(
    "FIG. 10 — Operator dashboard",
    "Screenshot of the Firebase-driven dashboard during an active S3 "
    "ALARM: density-index gauge, FSM state badge, live frame, valve "
    "state, captured-mass reading, and audit event log.",
    "Fig. 10. Operator dashboard rendered from the Firebase Realtime "
    "Database during an active S3 ALARM event."
)

add_subsection_heading("G.  System-Level Trade-offs and Limitations")
add_para(
    "Three trade-offs warrant discussion. First, the 2.5 s S3 "
    "confirmation window is the single largest contributor to end-to-"
    "end latency; shortening it would improve responsiveness at the "
    "cost of higher false-alarm rates and unnecessary valve closures, "
    "and high-throughput industrial channels may justify a shorter "
    "window paired with a stricter AI threshold. Second, the 0.60 "
    "mAP@50 on fragments is the weakest classifier output; in "
    "source-level industrial deployment fresh nurdles dominate the "
    "contamination profile and are handled well (mAP@50 ≥ 0.75), "
    "but aged or fragmented samples — dominant in ambient coastal "
    "water — would require a larger weathered-sample training set or "
    "an upgraded YOLOv8s/m model on a Jetson Orin Nano. Third, the "
    "current prototype is a single inline unit; multi-channel "
    "industrial bays will require either multiple units or a manifold, "
    "and the Firebase schema’s bay_id field already anticipates "
    "this topology."
)

add_subsection_heading("H.  Comparison with Existing Approaches")
add_para(
    "Table I contrasts NurdleDNA with conventional monitoring "
    "approaches. FTIR and Raman spectroscopy offer high chemical "
    "specificity but operate offline. Turbidity sensors and particle "
    "counters cannot reliably classify nurdles versus bubbles, sand, or "
    "biological particles. Passive filtration only records "
    "contamination, whereas NurdleDNA both records (via the cloud audit "
    "channel) and physically arrests it via the servo pinch valve."
)
add_figure(
    "table2_comparison.png",
    "Table I. Comparison of NurdleDNA with conventional and lab-grade "
    "microplastic-monitoring approaches across six operational criteria.",
    width_inches=3.4,
)

# =====================================================================
# V. Conclusion and Future Work  (compressed)
# =====================================================================
add_section_heading("V.  Conclusion and Future Work")
add_para(
    "This paper presented NurdleDNA, an edge-AI and IoT-enabled inline "
    "system that addresses a specific gap in industrial water "
    "monitoring: the absence of real-time, source-level interception of "
    "pre-production plastic pellets. By integrating a YOLOv8n vision "
    "pipeline on an NVIDIA Jetson Nano with a deterministic Moore FSM "
    "on Arduino, multi-sensor fusion across optical, gas, and "
    "gravimetric channels, and a Firebase-backed cloud audit log, the "
    "system links detection directly to physical containment and "
    "accountable reporting. The system was demonstrated live at the "
    "AURAK Excellence Award showcase on 21 May 2026, where it qualified "
    "for Tier 1 in the Sustainability, Economy and Environment "
    "category — the sole qualifying entry from the University of "
    "Wollongong in Dubai — and received letters of appreciation from "
    "industry partners BCL and Prismix."
)
add_para(
    "Three directions are identified for future work. (1) "
    "Field-grade hardening and multi-bay deployment, requiring IP-"
    "rated housing and SCADA integration; the Firebase bay_id schema "
    "already anticipates this. (2) Improved fragment-class "
    "classification via a larger weathered-sample training set and a "
    "YOLOv8s/m backbone on a Jetson Orin Nano. (3) Polymer-type "
    "identification through compact Raman or VIS-NIR spectroscopy at "
    "the evidence-cartridge stage — closing the loop on the conceptual "
    "“Nurdle DNA” framing and enabling source-traceable "
    "chain-of-custody records."
)

# =====================================================================
# References  (still in 2-column, smaller font)
# =====================================================================
add_section_heading("References")

REFS = [
    ("[1]",  "NOAA National Ocean Service, “What are microplastics?” "
             "National Oceanic and Atmospheric Administration, 2024. "
             "https://oceanservice.noaa.gov/facts/microplastics.html"),
    ("[2]",  "T. Ali, M. Mortula, B. Mohsen, L. Dronjak, R. Gawai, S. Atabay, "
             "Z. Khan, and K. Fattah, “Evaluating microplastic pollution "
             "along the Dubai coast: An empirical model combining on-site "
             "sampling and Sentinel-2 remote sensing data,” J. Sustain. "
             "Dev. Energy, Water Environ. Syst., vol. 12, no. 1, p. 1110482, "
             "2024."),
    ("[3]",  "S. Pu, H. Bushnaq, C. Munro, Y. Gibert, R. Sharma, V. Mishra, and "
             "L. F. Dumée, “Perspectives on transport pathways of "
             "microplastics across the Middle East and North Africa (MENA) "
             "region,” npj Clean Water, vol. 7, art. 114, 2024."),
    ("[4]",  "United Nations Environment Programme, “Plastic pollution,"
             "” UNEP, 2025. https://www.unep.org/plastic-pollution"),
    ("[5]",  "G. Sainger, “Microplastic Pollution in Oceans: A Barrier to "
             "Achieve Low Carbon Society,” IOP Conf. Ser.: Earth Environ. "
             "Sci., vol. 1279, art. 012021, 2023."),
    ("[6]",  "Z. Huang, B. Hu, and H. Wang, “Analytical methods for "
             "microplastics in the environment: A review,” Environ. Chem. "
             "Lett., vol. 21, no. 1, pp. 383–401, 2023."),
    ("[7]",  "J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, “You "
             "Only Look Once: Unified, real-time object detection,” in "
             "Proc. IEEE CVPR, 2016, pp. 779–788."),
    ("[8]",  "NVIDIA, “NVIDIA Jetson embedded systems developer kits and "
             "modules,” 2026. "
             "https://www.nvidia.com/en-us/autonomous-machines/embedded-systems/"),
    ("[9]",  "A. Donato et al., “Monitoring plastic pellet pollution in "
             "coastal environments through handheld Raman spectroscopy: Data "
             "from the Mediterranean coasts (Southern Italy),” Mar. "
             "Pollut. Bull., vol. 202, art. 116312, 2024."),
    ("[10]", "R.-D. Kulko, A. Pletl, A. Hanus, and B. Elser, “Detection "
             "of plastic granules and their mixtures,” Sensors, vol. 23, "
             "no. 7, art. 3441, 2023."),
    ("[11]", "T. Islam, M. R. Repon, M. A. Alim, T. Islam, and S. Shukhratov, "
             "“Detection and characterisation techniques for microfiber "
             "in wastewater,” in Microfiber Pollution, CRC Press, 2024, "
             "pp. 63–81."),
    ("[12]", "I. Shukhratov et al., “Optical detection of plastic waste "
             "through computer vision,” Intell. Syst. Appl., vol. 22, "
             "p. 200341, 2024."),
    ("[13]", "NVIDIA Developer, “Get started with Jetson Nano Developer "
             "Kit,” 2026. "
             "https://developer.nvidia.com/embedded/learn/get-started-jetson-nano-devkit"),
    ("[14]", "G. Jocher et al., “Ultralytics YOLOv8,” 2023. "
             "https://github.com/ultralytics/ultralytics"),
    ("[15]", "National Institute of Standards and Technology, “Finite "
             "state machine,” Dictionary of Algorithms and Data "
             "Structures. "
             "https://xlinux.nist.gov/dads/HTML/finite.html"),
]

for tag, body in REFS:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(tag + "  ")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(8.5)
    r2 = p.add_run(body)
    r2.font.name = "Times New Roman"; r2.font.size = Pt(8.5)

# =====================================================================
doc.save(OUT_PATH)
print("Saved:", OUT_PATH)
